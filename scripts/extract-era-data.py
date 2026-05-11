#!/usr/bin/env python3
"""Extract dynasty/ruler/era data from GB/T XXXXX.2 docx into structured JSON."""
import json, sys, re
from docx import Document

DOCX = '/Users/mulgogi/src/chinese/library/reference-docs/《日期和时间 日历体系代码 第2部分：中国古代》-征求意见稿-20250627.docx'

doc = Document(DOCX)
table = doc.tables[5]  # 历代纪元表 (650 rows)

def parse_year(s):
    """Parse year string like '公元618' or '公元前206' or '' into int (negative=BCE)"""
    if not s:
        return None
    s = s.strip()
    m = re.match(r'公元前(\d+)', s)
    if m:
        return -int(m.group(1))
    m = re.match(r'公元(\d+)', s)
    if m:
        return int(m.group(1))
    return None

# Extract all entries with cell inheritance
keys = ['dynasty','dynasty_code','ruler','ruler_code','reign_count','era_name','era_code','era_instance','ganzhi','gregorian']
entries = []
last = {}
for row in table.rows[1:]:
    cells = [cell.text.strip() for cell in row.cells]
    if cells[0].startswith('全表参考'):
        break
    entry = {}
    for j, key in enumerate(keys):
        val = cells[j] if j < len(cells) else ''
        if val:
            entry[key] = val
            last[key] = val
        else:
            entry[key] = last.get(key, '')
    entry['year'] = parse_year(entry.get('gregorian', ''))
    entries.append(entry)

# ─── Build dynasty table ────────────────────────────────────────
dynasties = []
seen_dyn = set()
for e in entries:
    code = e['dynasty_code']
    if code not in seen_dyn:
        seen_dyn.add(code)
        dyn_years = [e2['year'] for e2 in entries if e2['dynasty_code'] == code and e2['year'] is not None]
        dynasties.append({
            'code': int(code),
            'label': e['dynasty'],
            'firstYear': min(dyn_years) if dyn_years else None,
            'lastYear': max(dyn_years) if dyn_years else None,
            'rulerCount': len(set((e2['ruler_code'], e2['reign_count']) for e2 in entries if e2['dynasty_code'] == code)),
        })

# ─── Build ruler table ──────────────────────────────────────────
rulers = []
seen_ruler = set()
for e in entries:
    key = (e['dynasty_code'], e['ruler_code'], e['reign_count'])
    if key not in seen_ruler:
        seen_ruler.add(key)
        ruler_years = [e2['year'] for e2 in entries
                       if e2['dynasty_code'] == e['dynasty_code']
                       and e2['ruler_code'] == e['ruler_code']
                       and e2['reign_count'] == e['reign_count']
                       and e2['year'] is not None]
        rulers.append({
            'dynastyCode': int(e['dynasty_code']),
            'rulerCode': int(e['ruler_code']),
            'reignCount': int(e['reign_count']),
            'label': e['ruler'],
            'firstYear': min(ruler_years) if ruler_years else None,
            'lastYear': max(ruler_years) if ruler_years else None,
        })

# ─── Build era name table ───────────────────────────────────────
eras = []
seen_era = set()
for e in entries:
    if not e.get('era_name'):
        continue
    key = (e['dynasty_code'], e['ruler_code'], e['reign_count'], e['era_code'], e['era_instance'])
    if key not in seen_era:
        seen_era.add(key)
        era_years = [e2['year'] for e2 in entries
                     if e2['dynasty_code'] == e['dynasty_code']
                     and e2['ruler_code'] == e['ruler_code']
                     and e2['reign_count'] == e['reign_count']
                     and e2.get('era_code') == e['era_code']
                     and e2.get('era_instance') == e['era_instance']
                     and e2['year'] is not None]
        eras.append({
            'dynastyCode': int(e['dynasty_code']),
            'rulerCode': int(e['ruler_code']),
            'reignCount': int(e['reign_count']),
            'eraCode': int(e['era_code']),
            'eraInstance': int(e['era_instance']),
            'label': e['era_name'],
            'firstYear': min(era_years) if era_years else None,
            'lastYear': max(era_years) if era_years else None,
        })

# ─── Build ganzhi table ─────────────────────────────────────────
ganzhi_table = doc.tables[12]
ganzhi_entries = []
for row in ganzhi_table.rows[1:]:
    cells = [cell.text.strip() for cell in row.cells]
    if not cells[0] or not cells[0].startswith('公元'):
        continue
    ganzhi_entries.append({
        'year': parse_year(cells[0]),
        'ganzhi': cells[1],
        'ganzhiCode': int(cells[2]),
        'calendarCode': cells[3],
    })

# ─── Output ─────────────────────────────────────────────────────
output = {
    'dynasties': dynasties,
    'rulers': rulers,
    'eras': eras,
    'ganzhi': ganzhi_entries,
}

print(json.dumps(output, ensure_ascii=False, indent=2))
print(f'# Stats: {len(dynasties)} dynasties, {len(rulers)} rulers, {len(eras)} era names, {len(ganzhi_entries)} ganzhi entries', file=sys.stderr)
